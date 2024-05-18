from spacetype import Room, Apartment, MultistoryBuilding
import docx
from kivy.app import App
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.dropdown import DropDown
from kivy.uix.textinput import TextInput
from kivy.uix.boxlayout import BoxLayout
from kivy.core.window import Window

class MyApp(App):
    def __init__(self):
        super().__init__()
        Window.size = (400, 350)
        Window.clearcolor = (200/255, 150/255, 150/255, 1)
        self.title = "Калькулятор тепловой мощности"

        self.dropdown = DropDown()
        self.mainbutton = Button(text='Комната')
        self.space_label = Label(text='Выберите тип помещения:')
        self.lenght_imput = TextInput(hint_text='Длина комнаты (м)', multiline=False)
        self.width_imput = TextInput(hint_text='Ширина комнаты (м)', multiline=False)
        self.height_imput = TextInput(hint_text = 'Высота комнаты (м)', multiline=False)
        self.room_num_imput = TextInput(hint_text='Кол-во комнат (для квартиры/многоэтажного дома)', multiline=False)
        self.floor_num_imput = TextInput(hint_text="Кол-во этажей (для многоэтажного дома)", multiline=False)
        self.calc_button = Button(text='Расчитать и сохранить')
        self.calc_button.bind(on_release=self.calc_result)
        self.result_label = Label(text='Результат', size_hint=(1, 1.5))

    def build(self):
        box = BoxLayout(orientation='vertical', spacing=3)

        box.add_widget(self.space_label)

        for space in ("Комната", "Квартира", "Многоэтажный дом"):
            btn = Button(text=space, size_hint_y=None, height=44)
            btn.bind(on_release=lambda btn: self.dropdown.select(btn.text))
            self.dropdown.add_widget(btn)
        self.mainbutton.bind(on_release=self.dropdown.open)
        self.dropdown.bind(on_select=lambda instance, x: setattr(self.mainbutton, 'text', x))

        box.add_widget(self.mainbutton)
        box.add_widget(self.lenght_imput)
        box.add_widget(self.width_imput)
        box.add_widget(self.height_imput)
        box.add_widget(self.room_num_imput)
        box.add_widget(self.floor_num_imput)
        box.add_widget(self.calc_button)
        box.add_widget(self.result_label)

        return box

    def calc_result(self, *args):
        if (self.lenght_imput.text != "") & (self.width_imput.text != "") & (self.height_imput.text != "") & (self.height_imput.text != "") & (self.height_imput.text != ""):
            lenght = float(self.lenght_imput.text)
            width = float(self.width_imput.text)
            height = float(self.height_imput.text)

            if self.mainbutton.text == "Комната":
                building = Room(height, width, lenght)
            elif self.mainbutton.text == "Квартира":
                room_num = float(self.room_num_imput.text)
                building = Apartment(height, width, lenght, room_num)
            else:
                room_num = float(self.room_num_imput.text)
                floor_num = float(self.floor_num_imput.text)
                building = MultistoryBuilding(height, width, lenght, floor_num ,room_num)
                
            self.result_label.text = f"{str(building)}\nРезультат сохранен в docx файл"

            doc = docx.Document()
            doc.add_heading('Результаты расчетов', level=1)
            doc.add_paragraph(f"{str(building)}")
            doc.save('report.docx')
        
if __name__ == "__main__":
	MyApp().run()